import requests
import re
import io
from docx import Document

def get_coi_data():
    """Získá data z webu České obchodní inspekce"""
    url = "https://www.coi.gov.cz/pro-spotrebitele/rizikove-e-shopy/"
    domains = set()
    try:
        headers = {'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'}
        r = requests.get(url, headers=headers, timeout=30)
        # Hledáme domény v buňkách tabulky <td>...</td>
        found = re.findall(r'<td>([a-z0-9.-]+\.[a-z]{2,10})</td>', r.text.lower())
        for d in found:
            clean = d.strip()
            if clean and '.' in clean:
                domains.add(clean)
    except Exception as e:
        print(f"Chyba u ČOI: {e}")
    return domains

def get_soi_data():
    """Stáhne a přečte DOCX soubor ze stránek slovenské SOI"""
    # Přímý odkaz na aktuální DOCX soubor
    url = "https://www.soi.sk/files/documents/rizikove-eshopy/zoznam-rizikovych-eshopov.docx"
    domains = set()
    try:
        headers = {'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'}
        r = requests.get(url, headers=headers, timeout=30)
        r.raise_for_status()
        
        # Načtení obsahu Wordu z paměti (BytesIO)
        doc = Document(io.BytesIO(r.content))
        full_text = []
        
        # Vytáhneme text z odstavců
        for para in doc.paragraphs:
            full_text.append(para.text)
        
        # Vytáhneme text ze všech tabulek v dokumentu
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    full_text.append(cell.text)
        
        # Spojíme vše do jednoho textu a hledáme domény
        combined_text = " ".join(full_text).lower()
        # Regulární výraz pro hledání domén (včetně těch s www.)
        found = re.findall(r'(?:www\.)?([a-z0-9.-]+\.[a-z]{2,10})', combined_text)
        
        for d in found:
            # Vyfiltrujeme systémový balast
            blacklist_keywords = ['soi.sk', 'gov.sk', 'microsoft', 'openxml', 'schema', 'w3.org', 'google', 'facebook']
            if d and '.' in d and not any(word in d for word in blacklist_keywords):
                domains.add(d.strip().strip('.'))
                
    except Exception as e:
        print(f"Chyba u SOI (DOCX): {e}")
    return domains

if __name__ == "__main__":
    # 1. Získání dat z obou zdrojů
    cz_domains = get_coi_data()
    sk_domains = get_soi_data()
    
    # 2. Spojení do jedné unikátní sady
    all_domains = cz_domains | sk_domains
    
    # 3. Formátování pro AdGuard/uBlock (||domena.cz^) a seřazení
    # Odfiltrujeme příliš krátké řetězce, které nejsou doménami
    final_list = sorted([f"||{d}^" for d in all_domains if len(d) > 3])

    # 4. Uložení do souboru blocklist.txt
    if final_list:
        with open("blocklist.txt", "w", encoding="utf-8") as f:
            f.write("! Title: Oficiální CZ & SK Blocklist (COI + SOI DOCX)\n")
            f.write(f"! Total items: {len(final_list)}\n")
            f.write("! Expires: 1 days\n")
            f.write("! --------------------------------------------------\n\n")
            f.write("\n".join(final_list))
        print(f"Úspěšně vytvořeno! Celkem nalezeno {len(final_list)} domén.")
    else:
        print("Nebyla nalezena žádná data, soubor nebyl aktualizován.")
